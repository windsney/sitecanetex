from django.shortcuts import render,reverse
from .models import Sindicancia,Sindicado,Ofendido,Testemunha,Oficio
from django.views.generic import TemplateView,ListView,DetailView, FormView
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import locale
from django.shortcuts import render, get_object_or_404, redirect
from django.conf import settings
from .forms import SindicadoForm,SindicanciaForm,TestemunhaForm,OfendidoForm,UsuarioForm,NotificarTestForm,NotificarOfenForm,NotificarSindForm,PrazoForm,Oficio_diversoForm,JuntadaaForm
from datetime import datetime
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .models import Sindicancia
import os
from django.conf import settings
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Cm,Pt,RGBColor

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate
from django.http import HttpResponse
from meu_modulo.data_escrita import *
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.decorators import login_required
from django.db.models import Max






'''class Home(TemplateView):
    template_name = "login.html"
'''

class Sind_Cadastradas(LoginRequiredMixin,ListView):

    template_name = "sindicancias_cadastradas.html"
    model = Sindicancia

    def get_queryset(self):
        # Filtra as sindicâncias para exibir apenas as cadastradas pelo usuário logado
        return Sindicancia.objects.filter(usuario=self.request.user.id)

'''class Detalhe_sind(DetailView):

    template_name = "detalhe_sind.html"
    model = Sindicancia
'''
'''class Cadastro_investigado_sind(DetailView):

    template_name = "cad_investigado_sind.html"
    model = Sindicado
'''
class Cd(TemplateView):

    template_name = "cad_cd.html"

class Inquerito (TemplateView):
    template_name = "cad_ipm.html"


class RIOG(TemplateView):
    template_name = "cad_riog.html"
class Criar_conta(FormView):
    template_name = "criar_conta.html"
    form_class = UsuarioForm
    def form_valid(self, form):
        form.save()
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('inicio:login')


@login_required(login_url='/')
def criar_sindicancia(request):
    usuario= request.user.id
    if request.method == 'POST':
        form = SindicanciaForm(request.POST)
        if form.is_valid():
            sindicancia = form.save(commit=False)
            sindicancia.usuario = request.user.id  # Atribui o valor de id do usuário ao objeto
            sindicancia.delegada = request.user.nome_completo
            sindicancia.unidade= request.user.unidade
            sindicancia.posto_delegada= request.user.posto
            sindicancia.rg_delegada= request.user.rgpm
            sindicancia.lotacao_delegada= request.user.unidade
            sindicancia.rua_quartel= request.user.rua
            sindicancia.numero_quartel= request.user.numero
            sindicancia.bairro_quartel= request.user.bairro
            sindicancia.email_quartel= request.user.email_bpm
            sindicancia.telefone_quartel= request.user.telefone
            sindicancia.cep_quartel= request.user.cep
            sindicancia.cidade_quartel= request.user.cidade
            sindicancia.save()
            form.save()
            return redirect('inicio:sind_cadastradas')  # Redirecione para uma página de sucesso após salvar
    else:
        form = SindicanciaForm()
    return render(request, 'criar_sindicancia.html', {'form': form,'usuario':usuario})

@login_required(login_url='/')
def editar_sindicancia(request, id):
    usuario = request.user.id
    sindicancia = get_object_or_404(Sindicancia, id=id)

    if request.method == 'POST':
        form = SindicanciaForm(request.POST, instance=sindicancia)
        if form.is_valid():
            form.save()
            return redirect('inicio:sind_cadastradas')  # Redirecione para uma página de sucesso após salvar
    else:
        form = SindicanciaForm(instance=sindicancia)

    print(sindicancia.data_portaria)

    return render(request, 'editar_sindicancia.html', {'form': form, 'sindicancia': sindicancia})

@login_required(login_url='/')
def excluir_sindicancia(request, id):
    sindicancia = get_object_or_404(Sindicancia, id=id)

    if request.method == 'POST':
        sindicancia.delete()
        return redirect('inicio:sind_cadastradas')  # Redirecionar após exclusão

    return render(request,'sindicancias_cadastradas.html' , {'sindicancia': sindicancia})
@login_required(login_url='/')
def detalhes_sindicancia(request, sindicancia_id):



    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    sindicados = sindicancia.sindicados.all()
    testemunhas = sindicancia.testemunhas.all()
    ofendidos = sindicancia.ofendidos.all()
    email = request.user.telefone


    # Adicione as queries para testemunhas, ofendidos e ofícios conforme necessário
    # testemunhas = sindicancia.testemunhas.all()
    # ofendidos = sindicancia.ofendidos.all()
    # oficios = sindicancia.oficios.all()
    vixi= len(sindicados)





    context = {
        'sindicancia': sindicancia,
        'sindicados': sindicados,
        'testemunhas': testemunhas,
        'ofendidos': ofendidos,
        'user_email': email,

        # Adicione as queries para testemunhas, ofendidos e ofícios conforme necessário
        # 'testemunhas': testemunhas,
        # 'ofendidos': ofendidos,

         'vixi': vixi,

    }


    return render(request, 'detalhe_sind.html', context)
@login_required(login_url='/')
def gerar_inicio_dos_trabalhos(request, sindicancia_id):
    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'inicio_dos_trabalhos.docx')

    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    sindicados = sindicancia.sindicados.all()
    usuario = request.user
    documento = DocxTemplate(template_path)  # DOCUMENTO  EXEMPLO#

    #dia_inicio = sindicancia.data_inicio

    mes_ini = sindicancia.data_inicio.month
    mesescritoini = mes_escrito(mes_ini)
    diaini = sindicancia.data_inicio.strftime('%d') if sindicancia.data_inicio else ''
    anoini = sindicancia.data_inicio.year

    dia_inicio = f'{diaini} de {mesescritoini} de {anoini}'


    posto_delegante = sindicancia.posto_delegante
    nome_delegante = sindicancia.delegante
    funcao_delegante = sindicancia.funcao_delegante
    portaria = sindicancia.numero

    #datada = sindicancia.data_portaria
    mes_port= sindicancia.data_portaria.month
    mesescritoport=mes_escrito(mes_port)
    diaport=sindicancia.data_portaria.strftime('%d') if sindicancia.data_portaria else ''
    anoport=sindicancia.data_portaria.year


    datada= f'{diaport} de {mesescritoport} de {anoport}'
    cidade = usuario.cidade
    unidade = usuario.unidade
    cr = usuario.cr
    unidade1 = usuario.unidade.upper()
    cr1 = usuario.cr.upper()
    rua = usuario.rua
    bairro = usuario.bairro
    numero = usuario.numero
    cep = usuario.cep
    email = usuario.email_bpm
    telefone= usuario.telefone




    #dia_recebido = sindicancia.dia_recebido
    mes_rec = sindicancia.dia_recebido.month
    mesescritorec = mes_escrito(mes_rec)
    diarec = sindicancia.dia_recebido.strftime('%d') if sindicancia.dia_recebido else ''
    anorec = sindicancia.dia_recebido.year

    dia_recebido = f'{diarec} de {mesescritorec} de {anorec}'



    nome_delegada = sindicancia.delegada
    posto_delegada = sindicancia.posto_delegada
    rg_delegada = sindicancia.rg_delegada
    context = {  # VARIÁ
        "dia_inicio": f"{dia_inicio}",
        "cr": f"{cr}",
        "cr1": f"{cr1}",
        "rua": f"{rua}",
        "bairro": f"{bairro}",
        "numero": f"{numero}",
        "telefone": f"{telefone}",
        "cep": f"{cep}",
        "email": f"{email}",
        "cidade": f"{cidade}",
        "unidade": f"{unidade}",
        "unidade1": f"{unidade1}",
        "posto_delegante": f"{posto_delegante}",
        "nome_delegante": f"{nome_delegante}",
        "funcao_delegante": f"{funcao_delegante}",
        "portaria": f"{portaria}",
        "datada": f"{datada}",
        "dia_recebido": f"{dia_recebido}",
        "nome_delegada": f"{nome_delegada}",
        "posto_delegada": f"{posto_delegada}",
        "rg_delegada": f"{rg_delegada}",

    }
    documento.render(context)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename= inicio_dos_trabalhos {portaria}.docx'
    documento.save(response)

    return response

@login_required(login_url='/')
def gerar_termo_abertura(request, sindicancia_id):
    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'termo_abertura.docx')

    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    sindicados = sindicancia.sindicados.all()
    usuario = request.user
    documento = DocxTemplate(template_path)  # DOCUMENTO  EXEMPLO#

    #dia_inicio = sindicancia.data_inicio

    mes_ini = sindicancia.data_inicio.month

    diaini = sindicancia.data_inicio.day
    anoini = sindicancia.data_inicio.year
    dia =dia_escrito(diaini)
    mes = mes_escrito(mes_ini)
    ano=ano_escrito(anoini)





    posto_delegante = sindicancia.posto_delegante
    nome_delegante = sindicancia.delegante
    funcao_delegante = sindicancia.funcao_delegante
    portaria = sindicancia.numero

    #datada = sindicancia.data_portaria
    mes_port= sindicancia.data_portaria.month
    mesescritoport=mes_escrito(mes_port)
    diaport=sindicancia.data_portaria.strftime('%d') if sindicancia.data_portaria else ''
    anoport=sindicancia.data_portaria.year


    datada= f'{diaport} de {mesescritoport} de {anoport}'




    #dia_recebido = sindicancia.dia_recebido
    mes_rec = sindicancia.dia_recebido.month
    mesescritorec = mes_escrito(mes_rec)
    diarec = sindicancia.dia_recebido.strftime('%d') if sindicancia.dia_recebido else ''
    anorec = sindicancia.dia_recebido.year

    dia_recebido = f'{diarec} de {mesescritorec} de {anorec}'

    cidade= usuario.cidade
    unidade=usuario.unidade
    cr = usuario.cr
    unidade1 = usuario.unidade.upper()
    cr1 = usuario.cr.upper()
    rua=usuario.rua
    bairro=usuario.bairro
    numero=usuario.numero
    telefone = usuario.telefone
    cep=usuario.cep
    email=usuario.email_bpm
    nome_delegada = usuario.nome_completo
    posto_delegada = usuario.posto
    rg_delegada = usuario.rgpm
    context = {  # VARIÁ
        "dia": f"{dia}",
        "mes": f"{mes}",
        "ano": f"{ano}",
        "cr": f"{cr}",
        "cr1": f"{cr1}",
        "rua": f"{rua}",
        "bairro": f"{bairro}",
        "numero": f"{numero}",
        "cep": f"{cep}",
        "email": f"{email}",
        "telefone": f"{telefone}",
        "cidade": f"{cidade}",
        "unidade": f"{unidade}",
        "unidade1": f"{unidade1}",
        "posto_delegante": f"{posto_delegante}",
        "nome_delegante": f"{nome_delegante}",
        "funcao_delegante": f"{funcao_delegante}",
        "portaria": f"{portaria}",
        "datada": f"{datada}",
        "dia_recebido": f"{dia_recebido}",
        "nome_delegada": f"{nome_delegada}",
        "posto_delegada": f"{posto_delegada}",
        "rg_delegada": f"{rg_delegada}",

    }
    documento.render(context)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename= Termo de Abertura {portaria}.docx'
    documento.save(response)

    return response



@login_required(login_url='/')
def gerar_declaracao_sindicado(request,sindicancia_id,id):
    sindicado = get_object_or_404(Sindicado, id=id)
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'declaracao_sindicado.docx')
    usuario = request.user





    documento = DocxTemplate(template_path)  # DOCUMENTO  EXEMPLO#

    #dia_inicio = sindicancia.data_inicio



    decla_sindicado = sindicado.declaracao
    nome_sindicado= sindicado.nome
    posto_sindicado= sindicado.posto_sindicado
    mae_sindicado= sindicado.mae
    pai_sindicado= sindicado.pai
    rg_sindicado= sindicado.rgpm
    lotacao_sindicado=sindicado.lotacao
    endereco_sindicado=sindicado.endereco
    portaria= sindicancia.numero

    mes_port = sindicancia.data_portaria.month
    mesescritoport = mes_escrito(mes_port)
    diaport = sindicancia.data_portaria.strftime('%d') if sindicancia.data_portaria else ''
    anoport = sindicancia.data_portaria.year

    datada = f'{diaport} de {mesescritoport} de {anoport}'

    hora_inicio=sindicado.hora_inicio
    hora_fim=sindicado.hora_fim
    sindicante= usuario.nome_completo
    posto_sindicante=usuario.posto
    rg_sindicante= usuario.rgpm
    cpf=sindicado.cpf
    email=sindicado.email
    telefone= sindicado.telefone

    mes_nasc = sindicado.data_nascimento.month
    mesescritonasc = mes_escrito(mes_nasc)
    dianasc = sindicado.data_nascimento.strftime('%d') if sindicancia.data_portaria else ''
    anonasc = sindicado.data_nascimento.year

    data_nascimento = f'{dianasc} de {mesescritonasc} de {anonasc}'



    mes_inqui = sindicado.data_inquiricao.month
    mesescritoinqui = mes_escrito(mes_inqui)
    diainqui = sindicado.data_inquiricao.day
    diaescritoinqui = dia_escrito(diainqui)
    anoinqui = sindicado.data_inquiricao.year
    anoescritoinqui = ano_escrito(anoinqui)

    data_inquiricao = f'{diaescritoinqui} dias do mês de {mesescritoinqui} do ano de {anoescritoinqui}'



    naturalidade=sindicado.naturalidade

    cr= usuario.cr.upper()

    lotacao_delegada= usuario.unidade
    lotacao_delegada1 = usuario.unidade.upper()
    rua_quartel=usuario.rua
    numero_quartel=usuario.numero
    bairro_quartel=usuario.bairro
    cidade_quartel=usuario.cidade
    cep_quartel=usuario.cep
    telefone_quartel=usuario.telefone
    email_quartel=usuario.email_bpm


    context = {  # VARIÁ
        "cr": f"{cr}",
        "lotacao_delegada": f"{lotacao_delegada}",
        "lotacao_delegada1": f"{lotacao_delegada1}",
        "rua_quartel": f"{rua_quartel}",
        "numero_quartel": f"{numero_quartel}",
        "bairro_quartel": f"{bairro_quartel}",
        "cidade_quartel": f"{cidade_quartel}",
        "cep_quartel": f"{cep_quartel}",
        "telefone_quartel": f"{telefone_quartel}",
        "email_quartel": f"{email_quartel}",

        "decla_sindicado": f"{decla_sindicado}",
        "nome_sindicado": f"{nome_sindicado}",
        "posto_sindicado": f"{posto_sindicado}",
        "mae_sindicado": f"{mae_sindicado}",
        "pai_sindicado": f"{pai_sindicado}",
        "rg_sindicado": f"{rg_sindicado}",
        "cpf": f"{cpf}",
        "email": f"{email}",
        "telefone": f"{telefone}",
        "data_nascimento": f"{data_nascimento}",
        "naturalidade": f"{naturalidade}",
       "endereco_sindicado": f"{endereco_sindicado}",
        "lotacao_sindicado": f"{lotacao_sindicado}",
        "portaria": f"{portaria}",
        "datada": f"{datada}",
        "hora_inicio": f"{hora_inicio}",
        "hora_fim": f"{hora_fim}",
        "sindicante": f"{sindicante}",
        "posto_sindicante": f"{posto_sindicante}",
        "rg_sindicante": f"{rg_sindicante}",
        "data_inquiricao": f"{data_inquiricao}"

    }
    documento.render(context)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename= Sindicado {nome_sindicado}.docx'
    documento.save(response)

    return response


@login_required(login_url='/')
def gerar_declaracao_testemunha(request,sindicancia_id,id):
    testemunha = get_object_or_404(Testemunha, id=id)
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    if testemunha.militar=='sim':
        template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'declaracao_testemunha_militar.docx')
        profissao = 'Policial Militar'


    else:
        profissao = testemunha.profissao



        template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'declaracao_testemunha.docx')
    usuario = request.user




    documento = DocxTemplate(template_path)  # DOCUMENTO  EXEMPLO#

    #dia_inicio = sindicancia.data_inicio



    decla_sindicado = testemunha.declaracao
    nome_sindicado= testemunha.nome
    graduacao = testemunha.graduacao

    mae_sindicado= testemunha.mae
    pai_sindicado= testemunha.pai
    rg_sindicado= testemunha.rgpm

    endereco_sindicado=testemunha.endereco
    portaria= sindicancia.numero

    mes_port = sindicancia.data_portaria.month
    mesescritoport = mes_escrito(mes_port)
    diaport = sindicancia.data_portaria.strftime('%d') if sindicancia.data_portaria else ''
    anoport = sindicancia.data_portaria.year

    datada = f'{diaport} de {mesescritoport} de {anoport}'

    hora_inicio=testemunha.hora_inicio
    hora_fim=testemunha.hora_fim
    sindicante= usuario.nome_completo
    posto_sindicante=usuario.posto
    rg_sindicante= usuario.rgpm
    cpf=testemunha.cpf
    email=testemunha.email
    telefone= testemunha.telefone

    mes_nasc = testemunha.data_nascimento.month
    mesescritonasc = mes_escrito(mes_nasc)
    dianasc = testemunha.data_nascimento.strftime('%d') if sindicancia.data_portaria else ''
    anonasc = testemunha.data_nascimento.year

    data_nascimento = f'{dianasc} de {mesescritonasc} de {anonasc}'



    mes_inqui = testemunha.data_inquiricao.month
    mesescritoinqui = mes_escrito(mes_inqui)
    diainqui = testemunha.data_inquiricao.day
    diaescritoinqui = dia_escrito(diainqui)
    anoinqui = testemunha.data_inquiricao.year
    anoescritoinqui = ano_escrito(anoinqui)

    data_inquiricao = f'{diaescritoinqui} dias do mês de {mesescritoinqui} do ano de {anoescritoinqui}'



    naturalidade=testemunha.naturalidade

    cr= usuario.cr.upper()

    lotacao_delegada= usuario.unidade
    lotacao_delegada1 = usuario.unidade.upper()
    rua_quartel=usuario.rua
    numero_quartel=usuario.numero
    bairro_quartel=usuario.bairro
    cidade_quartel=usuario.cidade
    cep_quartel=usuario.cep
    telefone_quartel=usuario.telefone
    email_quartel=usuario.email_bpm


    context = {  # VARIÁ
        "cr": f"{cr}",
        "lotacao_delegada": f"{lotacao_delegada}",
        "lotacao_delegada1": f"{lotacao_delegada1}",
        "rua_quartel": f"{rua_quartel}",
        "numero_quartel": f"{numero_quartel}",
        "bairro_quartel": f"{bairro_quartel}",
        "cidade_quartel": f"{cidade_quartel}",
        "cep_quartel": f"{cep_quartel}",
        "telefone_quartel": f"{telefone_quartel}",
        "email_quartel": f"{email_quartel}",

        "decla_sindicado": f"{decla_sindicado}",
        "nome_sindicado": f"{nome_sindicado}",

        "mae_sindicado": f"{mae_sindicado}",
        "pai_sindicado": f"{pai_sindicado}",
        "rg_sindicado": f"{rg_sindicado}",
        "cpf": f"{cpf}",
        "email": f"{email}",
        "telefone": f"{telefone}",
        "data_nascimento": f"{data_nascimento}",
        "naturalidade": f"{naturalidade}",
       "endereco_sindicado": f"{endereco_sindicado}",
        "profissao": f"{profissao}",
        "portaria": f"{portaria}",
        "datada": f"{datada}",
        "hora_inicio": f"{hora_inicio}",
        "hora_fim": f"{hora_fim}",
        "sindicante": f"{sindicante}",
        "posto_sindicante": f"{posto_sindicante}",
        "rg_sindicante": f"{rg_sindicante}",
        "data_inquiricao": f"{data_inquiricao}",
        "graduacao": f"{graduacao}"

    }
    documento.render(context)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename= Testemunha {nome_sindicado}.docx'
    documento.save(response)

    return response

@login_required(login_url='/')
def gerar_declaracao_ofendido(request, sindicancia_id, id):
    ofendido = get_object_or_404(Ofendido, id=id)
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'declaracao_ofendido.docx')

    documento = DocxTemplate(template_path)  # DOCUMENTO  EXEMPLO#

    # dia_inicio = sindicancia.data_inicio

    decla_sindicado = ofendido.declaracao
    nome_sindicado = ofendido.nome

    mae_sindicado = ofendido.mae
    pai_sindicado = ofendido.pai
    rg_sindicado = ofendido.rgpm
    profissao = ofendido.profissao
    endereco_sindicado = ofendido.endereco
    portaria = sindicancia.numero

    mes_port = sindicancia.data_portaria.month
    mesescritoport = mes_escrito(mes_port)
    diaport = sindicancia.data_portaria.strftime('%d') if sindicancia.data_portaria else ''
    anoport = sindicancia.data_portaria.year

    datada = f'{diaport} de {mesescritoport} de {anoport}'

    hora_inicio = ofendido.hora_inicio
    hora_fim = ofendido.hora_fim
    sindicante = sindicancia.delegada
    posto_sindicante = sindicancia.posto_delegada
    rg_sindicante = sindicancia.rg_delegada
    cpf = ofendido.cpf
    email = ofendido.email
    telefone = ofendido.telefone

    mes_nasc = ofendido.data_nascimento.month
    mesescritonasc = mes_escrito(mes_nasc)
    dianasc = ofendido.data_nascimento.strftime('%d') if sindicancia.data_portaria else ''
    anonasc = ofendido.data_nascimento.year

    data_nascimento = f'{dianasc} de {mesescritonasc} de {anonasc}'

    mes_inqui = ofendido.data_inquiricao.month
    mesescritoinqui = mes_escrito(mes_inqui)
    diainqui = ofendido.data_inquiricao.day
    diaescritoinqui = dia_escrito(diainqui)
    anoinqui = ofendido.data_inquiricao.year
    anoescritoinqui = ano_escrito(anoinqui)

    data_inquiricao = f'{diaescritoinqui} dias do mês de {mesescritoinqui} do ano de {anoescritoinqui}'

    naturalidade = ofendido.naturalidade

    cr = sindicancia.unidade.upper()

    lotacao_delegada = sindicancia.lotacao_delegada
    lotacao_delegada1 = sindicancia.lotacao_delegada.upper()
    rua_quartel = sindicancia.rua_quartel
    numero_quartel = sindicancia.numero_quartel
    bairro_quartel = sindicancia.bairro_quartel
    cidade_quartel = sindicancia.cidade_quartel
    cep_quartel = sindicancia.cep_quartel
    telefone_quartel = sindicancia.telefone_quartel
    email_quartel = sindicancia.email_quartel

    context = {  # VARIÁ
        "cr": f"{cr}",
        "lotacao_delegada": f"{lotacao_delegada}",
        "lotacao_delegada1": f"{lotacao_delegada1}",
        "rua_quartel": f"{rua_quartel}",
        "numero_quartel": f"{numero_quartel}",
        "bairro_quartel": f"{bairro_quartel}",
        "cidade_quartel": f"{cidade_quartel}",
        "cep_quartel": f"{cep_quartel}",
        "telefone_quartel": f"{telefone_quartel}",
        "email_quartel": f"{email_quartel}",

        "decla_sindicado": f"{decla_sindicado}",
        "nome_sindicado": f"{nome_sindicado}",

        "mae_sindicado": f"{mae_sindicado}",
        "pai_sindicado": f"{pai_sindicado}",
        "rg_sindicado": f"{rg_sindicado}",
        "cpf": f"{cpf}",
        "email": f"{email}",
        "telefone": f"{telefone}",
        "data_nascimento": f"{data_nascimento}",
        "naturalidade": f"{naturalidade}",
        "endereco_sindicado": f"{endereco_sindicado}",
        "profissao": f"{profissao}",
        "portaria": f"{portaria}",
        "datada": f"{datada}",
        "hora_inicio": f"{hora_inicio}",
        "hora_fim": f"{hora_fim}",
        "sindicante": f"{sindicante}",
        "posto_sindicante": f"{posto_sindicante}",
        "rg_sindicante": f"{rg_sindicante}",
        "data_inquiricao": f"{data_inquiricao}"

    }
    documento.render(context)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename= Ofendido {nome_sindicado}.docx'
    documento.save(response)

    return response



@login_required(login_url='/')
def gerar_relatorio(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)

    sindicados = Sindicado.objects.filter(portaria_id=sindicancia_id)
    ofendidos = Ofendido.objects.filter(portaria_id=sindicancia_id)
    usuario = request.user
    testemunhas = Testemunha.objects.filter(portaria_id=sindicancia_id)

#____________base de dados_______________________________________________


    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'relatorio_modelo.docx')
    # Criar um novo documento
    doc = Document(template_path)
    cr = usuario.cr.upper()
    unidade = usuario.unidade.upper()

    header_texts = [
        "POLÍCIA MILITAR DO ESTADO DE MATO GROSSO",
        f"{cr}",
        f"{unidade}"
    ]

    for text in header_texts:
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

# ____________________titulo cabeçalho___________________________________________________

    report_title = doc.add_paragraph("RELATÓRIO")
    report_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = report_title.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Adicionar espaçamento duplo antes do título "Relatório"
    report_title_format = report_title.paragraph_format
    report_title_format.space_after = Pt(24)  # Espaço duplo
    report_title_format.space_before = Pt(24)  # Espaço duplo

    #_____________________________ tituto relatorio___________________

    dados_paragraph = doc.add_paragraph()
    dados_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run1 = dados_paragraph.add_run("1. ")
    run2 = dados_paragraph.add_run("Dados")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = True
    run2.underline = True
    run1.bold = False

    #1. dados__________________________________________________________________
    portaria_paragraph = doc.add_paragraph()
    portaria_run = portaria_paragraph.add_run(f"Portaria nº: {sindicancia.numero.upper()}")
    portaria_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    portaria_run.font.name = 'Times New Roman'
    portaria_run.font.size = Pt(12)
    r = portaria_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


    if len(sindicados)==1: # verifica se tem um ou mais sindicados pra  colocar no plural se for o caso

        sindicado_paragraph = doc.add_paragraph()
        sindicado_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        sindicado_run = sindicado_paragraph.add_run("Sindicado: ")
        sindicado_run.font.name = 'Times New Roman'
        sindicado_run.font.size = Pt(12)
        sindicado_run.bold = True  # Deixar a palavra "Sindicado" em negrito
        r = sindicado_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        sindicado_paragraph = doc.add_paragraph()
        sindicado_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        sindicado_run = sindicado_paragraph.add_run("Sindicados: ")
        sindicado_run.font.name = 'Times New Roman'
        sindicado_run.font.size = Pt(12)
        sindicado_run.bold = False  # Deixar a palavra "Sindicado" em negrito
        r = sindicado_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for index, sindicado in enumerate(sindicados):
        if index == 0:
            sindicado_run = sindicado_paragraph.add_run(f"{sindicado.nome.upper()} - {sindicado.posto_sindicado.upper()}")
            sindicado_run.font.name = 'Times New Roman'
            sindicado_run.font.size = Pt(12)
            r = sindicado_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        else:
            sindicado_nome_paragraph = doc.add_paragraph()
            sindicado_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            sindicado_nome_paragraph.paragraph_format.left_indent = Pt(63)
            sindicado_nome_run = sindicado_nome_paragraph.add_run(f"{sindicado.nome.upper()} - {sindicado.posto_sindicado.upper()}")
            sindicado_nome_run.font.name = 'Times New Roman'
            sindicado_nome_run.font.size = Pt(12)
            r = sindicado_nome_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

#________________________ sindicados nome alinhado e se tiver mais de um colocar sindicados

    if len(ofendidos)==1: # verifica se tem um ou mais ofendido pra  colocar no plural se for o caso

        ofendido_paragraph = doc.add_paragraph()
        ofendido_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        ofendido_run = ofendido_paragraph.add_run("Ofendido: ")
        ofendido_run.font.name = 'Times New Roman'
        ofendido_run.font.size = Pt(12)
        ofendido_run.bold = False  # Deixar a palavra "Sindicado" em negrito
        r = ofendido_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        ofendido_paragraph = doc.add_paragraph()
        ofendido_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        ofendido_run = ofendido_paragraph.add_run("Ofendidos: ")
        ofendido_run.font.name = 'Times New Roman'
        ofendido_run.font.size = Pt(12)
        ofendido_run.bold = False  # Deixar a palavra "Sindicado" em negrito
        r = ofendido_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for index, ofendido in enumerate(ofendidos):
        if index == 0:
            ofendido_run = ofendido_paragraph.add_run(f"{ofendido.nome}")
            ofendido_run.font.name = 'Times New Roman'
            ofendido_run.font.size = Pt(12)
            r = ofendido_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        else:
            ofendido_nome_paragraph = doc.add_paragraph()
            ofendido_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ofendido_nome_paragraph.paragraph_format.left_indent = Pt(58)
            ofendido_nome_run = ofendido_nome_paragraph.add_run(f"{ofendido.nome}")
            ofendido_nome_run.font.name = 'Times New Roman'
            ofendido_nome_run.font.size = Pt(12)
            r = ofendido_nome_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    ###########_________ adicioona ofendido ou plural se necessário

    fato_paragraph = doc.add_paragraph()
    fato_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fato_paragraph.paragraph_format.first_line_indent = Pt(70)

    fato_run = fato_paragraph.add_run(f"Fato (s): {sindicancia.historico}")
    fato_run.font.name = 'Times New Roman'
    fato_run.font.size = Pt(12)
    r = fato_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    #________ coloca  o historico no  fato com espaço na primeira  linha  depois justificado.

    local_paragraph = doc.add_paragraph()
    local_run = local_paragraph.add_run(f"Local: {sindicancia.rua_fato}; Bairro {sindicancia.bairro_fato};Cidade {sindicancia.cidade_fato};")
    local_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    local_run.font.name = 'Times New Roman'
    local_run.font.size = Pt(12)
    r = local_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    data_hora_paragraph = doc.add_paragraph()
    data_formatada = sindicancia.data_fato.strftime('%d.%m.%Y')
    data_hora_run = data_hora_paragraph.add_run(f"Data/Hora: {data_formatada}/{sindicancia.hora_fato}")
    data_hora_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    data_hora_run.font.name = 'Times New Roman'
    data_hora_run.font.size = Pt(12)
    r = data_hora_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    em_servico_paragraph = doc.add_paragraph()
    em_servico_run = em_servico_paragraph.add_run(f"Em Serviço? {sindicancia.em_servico}")
    em_servico_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    em_servico_run.font.name = 'Times New Roman'
    em_servico_run.font.size = Pt(12)
    r = em_servico_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    testemunhas_ouvidas_paragraph = doc.add_paragraph()
    testemunhas_ouvidas_run = testemunhas_ouvidas_paragraph.add_run(f"Testemunhas Ouvidas:")
    testemunhas_ouvidas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    testemunhas_ouvidas_run.font.name = 'Times New Roman'
    testemunhas_ouvidas_run.font.size = Pt(12)
    r = testemunhas_ouvidas_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(testemunhas)==0:
        testemunha_nome_paragraph = doc.add_paragraph()
        testemunha_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        testemunha_nome_paragraph.paragraph_format.left_indent = Pt(30)
        testemunha_nome_run = testemunha_nome_paragraph.add_run(f"Não houveram testemunhas")
        testemunha_nome_run.font.name = 'Times New Roman'
        testemunha_nome_run.font.size = Pt(12)
        r = testemunha_nome_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    else:

        for testemunha in testemunhas:

            testemunha_nome_paragraph = doc.add_paragraph()
            testemunha_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            testemunha_nome_paragraph.paragraph_format.left_indent = Pt(30)
            testemunha_nome_run = testemunha_nome_paragraph.add_run(f"\u25CF    {testemunha.nome} - fls {testemunha.fls}")
            testemunha_nome_run.font.name = 'Times New Roman'
            testemunha_nome_run.font.size = Pt(12)
            r = testemunha_nome_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for sindicado in sindicados:
        sindicado_nome_paragraph = doc.add_paragraph()
        sindicado_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        #sindicado_nome_paragraph.paragraph_format.left_indent = Pt(63)
        sindicado_nome_run = sindicado_nome_paragraph.add_run(f" O Sindicado, {sindicado.nome.upper()} - {sindicado.posto_sindicado.upper()}, foi inquirido conforme fls: {sindicado.fls}")
        sindicado_nome_run.font.name = 'Times New Roman'
        sindicado_nome_run.font.size = Pt(12)
        r = sindicado_nome_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for ofendido in ofendidos:
        ofendido_nome_paragraph = doc.add_paragraph()
        ofendido_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        #sindicado_nome_paragraph.paragraph_format.left_indent = Pt(63)
        ofendido_nome_run = ofendido_nome_paragraph.add_run(f" O Ofendido, {ofendido.nome.upper()}, foi inquirido conforme fls:{ofendido.fls}")
        ofendido_nome_run.font.name = 'Times New Roman'
        ofendido_nome_run.font.size = Pt(12)
        r = ofendido_nome_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    objetos_apreendidos_paragraph = doc.add_paragraph()
    objetos_apreendidos_run = objetos_apreendidos_paragraph.add_run(f"Objeto(s) Apreendido(s):")
    objetos_apreendidos_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    objetos_apreendidos_run.font.name = 'Times New Roman'
    objetos_apreendidos_run.font.size = Pt(12)
    r = objetos_apreendidos_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    nao_houveram_paragraph = doc.add_paragraph()
    nao_houveram_run = nao_houveram_paragraph.add_run(f"\u25CF    Não Houve;")
    nao_houveram_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    nao_houveram_paragraph.paragraph_format.left_indent = Pt(30)
    nao_houveram_run.font.name = 'Times New Roman'
    nao_houveram_run.font.size = Pt(12)
    r = nao_houveram_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    diligencias_paragraph = doc.add_paragraph()
    diligencias_run = diligencias_paragraph.add_run(f"Diligências Realizadas:")
    diligencias_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    diligencias_run.font.name = 'Times New Roman'
    diligencias_run.font.size = Pt(12)
    r = diligencias_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f"\u25CF    Termo de Abertura ; Fls xx;")
    dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dili_paragraph.paragraph_format.left_indent = Pt(30)

    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f"\u25CF    Ofício de Início dos Trabalhos;  Fls xx;")
    dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dili_paragraph.paragraph_format.left_indent = Pt(30)
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(ofendidos)==0:
        pass
    else:
        for ofendido in ofendidos:

            dili_paragraph = doc.add_paragraph()
            dili_run = dili_paragraph.add_run(f"\u25CF    Termo de Perguntas ao Ofendido {ofendido.nome} Fls {ofendido.fls};")
            dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            dili_paragraph.paragraph_format.left_indent = Pt(30)
            dili_run.font.name = 'Times New Roman'
            dili_run.font.size = Pt(12)
            r = dili_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(testemunhas)==0:
        pass
    else:
        for testemunha in testemunhas:
            dili_paragraph = doc.add_paragraph()
            dili_run = dili_paragraph.add_run(f"\u25CF    Termo de Inquirição de Testemunha {testemunha.nome} Fls {testemunha.fls};")
            dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            dili_paragraph.paragraph_format.left_indent = Pt(30)
            dili_run.font.name = 'Times New Roman'
            dili_run.font.size = Pt(12)
            r = dili_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(sindicados)==0:
        pass
    else:
        for sindicado in sindicados:
            dili_paragraph = doc.add_paragraph()
            dili_run = dili_paragraph.add_run(f"\u25CF    Termo de Perguntas ao Sindicado {sindicado.nome} Fls {sindicado.fls};")
            dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            dili_paragraph.paragraph_format.left_indent = Pt(30)
            dili_run.font.name = 'Times New Roman'
            dili_run.font.size = Pt(12)
            r = dili_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dados_paragraph = doc.add_paragraph()
    dados_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run1 = dados_paragraph.add_run("2. ")
    run2 = dados_paragraph.add_run("Os Fatos")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = True
    run2.underline = True
    run1.bold = False

    # 2. Os Fatos__________________________________________________________________

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f"Do que foi apurado na presente sindicância constata-se que os fatos ocorreram da seguinte forma:")
    dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dili_paragraph.paragraph_format.left_indent = Pt(30)
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')



    fato_paragraph = doc.add_paragraph()
    fato_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fato_paragraph.paragraph_format.first_line_indent = Pt(40)

    fato_run = fato_paragraph.add_run(f"XXXXXXXXXcitar como ocorreuXXXXXXXXXXXX")
    fato_run.font.name = 'Times New Roman'
    fato_run.font.size = Pt(12)
    fato_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    r = fato_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
#________________________________ comentar fatos apurados_________________
    dados_paragraph = doc.add_paragraph()
    dados_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run1 = dados_paragraph.add_run("3. ")
    run2 = dados_paragraph.add_run("Análise dos Elementos Informativos")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = True
    run2.underline = True
    run1.bold = False
    # 3.	Análise dos Elementos Informativos__________________________________________________________________
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f"Diante da tentativa de esclarecer o evento, há necessidade de comentar, um a um, os depoimentos:")
    dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dili_paragraph.paragraph_format.left_indent = Pt(30)
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(ofendidos)==0:
        pass
    else:
        for ofendido in ofendidos:
            dili_paragraph = doc.add_paragraph()
            dili_run = dili_paragraph.add_run(f'\u25CF Diz o Ofendido, {ofendido.nome} Conforme Fls {ofendido.fls}')
            dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            dili_paragraph.paragraph_format.left_indent = Pt(70)
            dili_run.font.name = 'Times New Roman'
            dili_run.font.size = Pt(12)
            r = dili_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            dili_run2 = dili_paragraph.add_run(f' "{ofendido.declaracao}"')
            dili_run2.font.name = 'Times New Roman'
            dili_run2.font.size = Pt(12)
            dili_run2.italic = True
            r2 = dili_run2._element
            r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            dili_run2 = dili_paragraph.add_run('. Verificamos na declaração que xxxxxcomentário do encarregadoxxxxxx.')
            dili_run2.font.name = 'Times New Roman'
            dili_run2.font.size = Pt(12)
            dili_run2.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
            dili_run2.italic = False
            r2 = dili_run2._element
            r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


    if len(testemunhas)==0:
        pass
    else:
        for testemunha in testemunhas:
            dili_paragraph = doc.add_paragraph()
            dili_run = dili_paragraph.add_run(f'\u25CF Diz a Testemunha, {testemunha.nome} Conforme Fls {testemunha.fls}')
            dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            dili_paragraph.paragraph_format.left_indent = Pt(70)
            dili_run.font.name = 'Times New Roman'
            dili_run.font.size = Pt(12)
            r = dili_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            dili_run2 = dili_paragraph.add_run(f' "{testemunha.declaracao}"')
            dili_run2.font.name = 'Times New Roman'
            dili_run2.font.size = Pt(12)
            dili_run2.italic = True
            r2 = dili_run2._element
            r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            dili_run2 = dili_paragraph.add_run('. Verificamos na declaração que xxxxxcomentário do encarregadoxxxxxx.')
            dili_run2.font.name = 'Times New Roman'
            dili_run2.font.size = Pt(12)
            dili_run2.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
            dili_run2.italic = False
            r2 = dili_run2._element
            r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(sindicados)==0:
        pass
    else:
        for sindicado in sindicados:
            dili_paragraph = doc.add_paragraph()
            dili_run = dili_paragraph.add_run(f'\u25CF Diz o Sindicado, {sindicado.nome} Conforme Fls {sindicado.fls}')
            dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            dili_paragraph.paragraph_format.left_indent = Pt(70)
            dili_run.font.name = 'Times New Roman'
            dili_run.font.size = Pt(12)
            r = dili_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            dili_run2 = dili_paragraph.add_run(f' "{sindicado.declaracao}"')
            dili_run2.font.name = 'Times New Roman'
            dili_run2.font.size = Pt(12)
            dili_run2.italic = True
            r2 = dili_run2._element
            r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            dili_run2 = dili_paragraph.add_run('. Verificamos na declaração que xxxxxcomentário do encarregadoxxxxxx.')
            dili_run2.font.name = 'Times New Roman'
            dili_run2.font.size = Pt(12)
            dili_run2.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
            dili_run2.italic = False
            r2 = dili_run2._element
            r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# declaração e  comentario de cada envolvido cadastrado__________________________

    dados_paragraph = doc.add_paragraph()
    dados_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run1 = dados_paragraph.add_run("4. ")
    run2 = dados_paragraph.add_run("Da análise da defesa")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = True
    run2.underline = True
    run1.bold = False

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f" ###EXEMPLO APAGAR######### A defesa relata que ")
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dili_paragraph.paragraph_format.first_line_indent = Pt(30)

    #dili_paragraph.paragraph_format.left_indent = Pt(30)
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    dili_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dili_run2 = dili_paragraph.add_run(f' "houve inobservância a formalidades essenciais a regularidade do feito e sua ausência atenta contra o exercício da ampla defesa e o contraditório do Sindicado, pois a defesa diz que como teria ocasionado lesões físicas ao ofendido, entretanto não fora juntada nos autos o corpo de delito, conforme diz o  art. 328 do CPPM."')
    dili_run2.font.name = 'Times New Roman'
    dili_run2.font.size = Pt(12)
    dili_run2.font.color.rgb = RGBColor(0, 0, 255)  # Define a cor AZUL
    dili_run2.italic = True
    r2 = dili_run2._element
    r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dili_run2 = dili_paragraph.add_run(' No entanto o referido exame não foi produzido pela vítima tendo em vista que a mesma sequer fez a confecção de boletim de ocorrências, assim não foi possível realizar a juntada do exame de corpo de delito. E o objetivo do procedimento administrativo é verificar possíveis transgressões disciplinar cometida durante a abordagem.')
    dili_run2.font.name = 'Times New Roman'
    dili_run2.font.size = Pt(12)
    dili_run2.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    dili_run2.italic = False
    r2 = dili_run2._element
    r2.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


    # 4 analise da defesa__________________________________________________________________

    dados_paragraph = doc.add_paragraph()
    dados_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    #dados_paragraph.paragraph_format.first_line_indent = Pt(30)
    run1 = dados_paragraph.add_run("5. ")
    run2 = dados_paragraph.add_run("Solução")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = True
    run2.underline = True
    run1.bold = False

# 5 Solução__________________________________________________________________
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(
        f"Diante do exposto, após análise minuciosa dos autos, concluímos os trabalhos desta sindicância, de forma a perceber que:")
    dili_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dili_paragraph.paragraph_format.left_indent = Pt(30)
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    fato_paragraph = doc.add_paragraph()
    fato_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fato_paragraph.paragraph_format.first_line_indent = Pt(40)

    fato_run = fato_paragraph.add_run(f"XXXXXXXXXcitar como ocorreuXXXXXXXXXXXX")
    fato_run.font.name = 'Times New Roman'
    fato_run.font.size = Pt(12)
    fato_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    r = fato_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


    texto='Diante do que foi apurado, este Sindicante encontrou indícios de transgressões disciplinares com relação ao tapa desferido durante a abordagem que não condiz que o que preceitua o Manual de Procedimento Operacional Padrão da Polícia Militar do Estado de Mato Grosso (POP PMMT), pois apesar de a vítima não realizar a confecção do Boletim de Ocorrência o vídeo deixa clara a dinâmica da abordagem que culminou com o golpe na cabeça do Ofendido. Sobre a questão  de injuria racial/racismo, nada consta no  vídeo apresentado, citado pelo ofendido e pela testemunha que gravou o vídeo e  negado veemente por todos os policiais inquiridos.'
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(
        f"{texto}")
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dili_paragraph.paragraph_format.first_line_indent = Pt(30)

    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    dili_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if len(sindicados)==1:


        texto='Portanto salvo melhor juízo, há indícios de transgressão disciplinar; porém não de crime militar a ser imputado ao Sindicado:'
        sindicados_texto = ' '.join([f"{s.nome} - {s.posto_sindicado} RGPMMT {s.rgpm};" for s in sindicados])
        dili_paragraph = doc.add_paragraph()
        dili_run = dili_paragraph.add_run(f"{texto} {sindicados_texto}")

        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dili_paragraph.paragraph_format.first_line_indent = Pt(30)

        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
        r = dili_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        texto = 'Portanto salvo melhor juízo, há indícios de transgressão disciplinar; porém não de crime militar a serem imputados aos Sindicados:'
        sindicados_texto = ' '.join([f"{s.nome} - {s.posto_sindicado} RGPMMT {s.rgpm};" for s in sindicados])
        dili_paragraph = doc.add_paragraph()
        dili_run = dili_paragraph.add_run(f"{texto} {sindicados_texto}")

        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dili_paragraph.paragraph_format.first_line_indent = Pt(30)

        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
        r = dili_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
#___________________________relatório gerado ___________________________________________________________

    import locale

    # Definir a localidade para português do Brasil
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

    # Obter a data atual
    data_atual = datetime.now()

    # Formatar a data no formato desejado
    data_formatada = data_atual.strftime('%d de %B de %Y')

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run('')
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(
        f"Quartel do {usuario.unidade} de {usuario.cidade} {data_formatada}.")
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dili_paragraph.paragraph_format.first_line_indent = Pt(40)

    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    #dili_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


    tamanho= len(usuario.nome_completo)+len(usuario.posto)+6

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run('')
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run('_'*tamanho)
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    dili_paragraph.paragraph_format.space_before = Pt(0)
    dili_paragraph.paragraph_format.space_after = Pt(1)
    dili_paragraph.paragraph_format.line_spacing = Pt(12)




    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'{usuario.nome_completo} - {usuario.posto}')
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    dili_paragraph.paragraph_format.space_before = Pt(0)
    dili_paragraph.paragraph_format.space_after = Pt(1)
    dili_paragraph.paragraph_format.line_spacing = Pt(12)
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'RGPPM {usuario.rgpm} - Sindicante')
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)


    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    # dili_run.font.color.rgb = RGBColor(255, 0, 0)  # Define a cor vermelha
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    rodape(doc=doc, usuario=usuario)


    # Nome do arquivo
    nome = f'Relatorio sind{sindicancia.numero}'

    # Configurar a resposta HTTP para download do arquivo
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{nome}.docx"'

    # Salvar o documento modificado na resposta
    try:
        doc.save(response)
    except Exception as e:
        return HttpResponse(f"Erro ao salvar o documento: {e}", status=500)

    return response




@login_required(login_url='/')
def cadastrar_sindicado(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)

    if request.method == 'POST':
        form = SindicadoForm(request.POST)
        if form.is_valid():
            sindicado = form.save(commit=False)
            sindicado.portaria = sindicancia
            sindicado.save()
            return redirect('inicio:detalhes_sindicancia', sindicancia_id=sindicancia.id)
    else:
        form = SindicadoForm()

    return render(request, 'cadastrar_sindicado.html', {'form': form, 'sindicancia': sindicancia})

@login_required(login_url='/')
def editar_sindicado(request,sindicancia_id, id):
    sindicado = get_object_or_404(Sindicado, id=id)
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    if request.method == 'POST':
        form = SindicadoForm(request.POST, instance=sindicado)
        if form.is_valid():
            form.save()
            return redirect('inicio:detalhes_sindicancia',sindicancia_id=sindicancia.id)
    else:
        form = SindicadoForm(instance=sindicado)


    return render(request, 'editar_sindicado.html', {'form': form,'sindicado':sindicado,'sindicancia':sindicancia})
@login_required(login_url='/')
def excluir_sindicado(request, sindicancia_id, id):
    sindicado = get_object_or_404(Sindicado, id=id)
    if request.method == 'POST':
        sindicado.delete()
        return redirect('inicio:detalhes_sindicancia', sindicancia_id=sindicancia_id)
    return render(request, 'sindicancias_cadastradas.html', {'sindicado': sindicado})



@login_required(login_url='/')
def cadastrar_testemunha(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)

    if request.method == 'POST':
        form = TestemunhaForm(request.POST)
        if form.is_valid():
            testemunha = form.save(commit=False)
            testemunha.portaria = sindicancia
            testemunha.save()
            return redirect('inicio:detalhes_sindicancia', sindicancia_id=sindicancia.id)
    else:
        form = TestemunhaForm()

    return render(request, 'cadastrar_testemunha.html', {'form': form, 'sindicancia': sindicancia})

@login_required(login_url='/')
def editar_testemunha(request,sindicancia_id, id):
    testemunha = get_object_or_404(Testemunha, id=id)
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    if request.method == 'POST':
        form = TestemunhaForm(request.POST, instance=testemunha)
        if form.is_valid():
            form.save()
            return redirect('inicio:detalhes_sindicancia',sindicancia_id=sindicancia.id)
    else:
        form = TestemunhaForm(instance=testemunha)


    return render(request, 'editar_testemunha.html', {'form': form,'testemunha':testemunha,'sindicancia':sindicancia})
@login_required(login_url='/')
def excluir_testemunha(request, sindicancia_id, id):
    testemunha = get_object_or_404(Testemunha, id=id)
    if request.method == 'POST':
        testemunha.delete()
        return redirect('inicio:detalhes_sindicancia', sindicancia_id=sindicancia_id)
    return render(request, 'sindicancias_cadastradas.html', {'testemunha': testemunha})

@login_required(login_url='/')
def cadastrar_ofendido(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)

    if request.method == 'POST':
        form = OfendidoForm(request.POST)
        if form.is_valid():
            ofendido = form.save(commit=False)
            ofendido.portaria = sindicancia
            ofendido.save()
            return redirect('inicio:detalhes_sindicancia', sindicancia_id=sindicancia.id)
    else:
        form = OfendidoForm()

    return render(request, 'cadastrar_ofendido.html', {'form': form, 'sindicancia': sindicancia})

@login_required(login_url='/')
def editar_ofendido(request,sindicancia_id, id):
    ofendido = get_object_or_404(Ofendido, id=id)
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    if request.method == 'POST':
        form = OfendidoForm(request.POST, instance=ofendido)
        if form.is_valid():
            form.save()
            return redirect('inicio:detalhes_sindicancia',sindicancia_id=sindicancia.id)
    else:
        form = OfendidoForm(instance=ofendido)


    return render(request, 'editar_ofendido.html', {'form': form,'ofendido':ofendido,'sindicancia':sindicancia})
@login_required(login_url='/')
def excluir_ofendido(request, sindicancia_id, id):
    ofendido = get_object_or_404(Ofendido, id=id)
    if request.method == 'POST':
        ofendido.delete()
        return redirect('inicio:detalhes_sindicancia', sindicancia_id=sindicancia_id)
    return render(request, 'sindicancias_cadastradas.html', {'ofendido': ofendido})

@login_required(login_url='/')
def gerar_remessa_dos_autos(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)

    sindicados = Sindicado.objects.filter(portaria_id=sindicancia_id)
    ofendidos = Ofendido.objects.filter(portaria_id=sindicancia_id)
    usuario = request.user
    testemunhas = Testemunha.objects.filter(portaria_id=sindicancia_id)
    padrao=sindicancia.padrao_oficio
    ano_corrente = datetime.now().year




    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'relatorio_modelo.docx')
    # Criar um novo documento
    doc = Document(template_path)
    oficio = Oficio()
    oficio.numero=numero_oficio(sindicancia_id)
    oficio.tipo='Ofício de Remessa dos Autos'

















    cabecalho(doc,usuario)
    criar_paragrafo(doc,'')

    #titulo_oficio(doc, usuario, tipo, sindicancia_id)


    oficio.id_portaria=sindicancia_id
    oficio.numero=numero_oficio(sindicancia_id) #oficio é a instancia do model que tem a coluna numero /  ja o numero_oficio é funçao para atribuir  numero sequancial soma 1 no maior numero da tabela


    titulo_oficio(doc, usuario, sindicancia_id)
    ###coloca numero no oficio




    char_style = doc.styles.add_style('CustomCharStyle', WD_STYLE_TYPE.CHARACTER)
    char_style.font.name = 'Times New Roman'
    char_style.font.size = Pt(12)
    char_style.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor preta
    r = char_style.element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Parágrafo 1
    datata = datetime.today()
    dia = f"{datata.day:02d}"
    mes = datata.month
    ano = datata.year

    paragraph1 = doc.add_paragraph()

    paragraph1.paragraph_format.space_before = Pt(15)


    paragraph1 = doc.add_paragraph()

    run1 = paragraph1.add_run('Ao Senhor Ten Cel PM João das Dores.')
    run1.style = 'CustomCharStyle'
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph1.paragraph_format.space_before = Pt(0)
    paragraph1.paragraph_format.space_after = Pt(0)

    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run(f'{sindicancia.funcao_delegante}.')
    run1.style = 'CustomCharStyle'
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph1.paragraph_format.space_before = Pt(0)
    paragraph1.paragraph_format.space_after = Pt(0)

    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run('Assunto: Remessa de Sindicância.')
    run1.style = 'CustomCharStyle'
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph1.paragraph_format.space_before = Pt(0)
    paragraph1.paragraph_format.space_after = Pt(0)

    paragraph1 = doc.add_paragraph()
    data_original = f'{sindicancia.data_portaria}'

    # Converter a string para um objeto datetime
    data_obj = datetime.strptime(data_original, '%Y-%m-%d')
    # Formatar a data no novo formato
    data_formatada = data_obj.strftime('%d.%m.%y')
    run1 = paragraph1.add_run(f'Ref: Portaria nº{sindicancia.numero}, datada de {data_formatada}.')
    run1.style = 'CustomCharStyle'
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph1.paragraph_format.space_before = Pt(0)
    paragraph1.paragraph_format.space_after = Pt(0)

    paragraph1 = doc.add_paragraph()
    paragraph1 = doc.add_paragraph('')
    paragraph1.paragraph_format.first_line_indent = Pt(120)
    run1 = paragraph1.add_run('Senhor Comandante,')
    run1.style = 'CustomCharStyle'
    paragraph1.paragraph_format.space_before = Pt(15)
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if len(sindicados) == 1:

        datata= sindicancia.data_portaria
        dia = f"{datata.day:02d}"
        mes = datata.month
        ano = datata.year


        texto = f'Tendo este Sindicante sido designado por Vossa Senhoria, e por motivo ter  finalizado os trabalhos referentes à Sindicância, instaurada pela Portaria nº {sindicancia.numero}  de {dia} de {mes_escrito(mes)} de {ano}, qual figura como Sindicado o'
        texto1 = 'estão conclusos  remeto a V.S.ª estes autos com xxxxxxxxxxxxxxxx folhas, para fins de Solução.'
        sindicados_texto = ' '.join([f" {s.posto_sindicado} {s.nome}, RG {s.rgpm} PMMT," for s in sindicados])
        dili_paragraph = doc.add_paragraph()
        dili_run = dili_paragraph.add_run(f"{texto} {sindicados_texto} {texto1}")

        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dili_paragraph.paragraph_format.first_line_indent = Pt(30)

        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_run.font.color.rgb = RGBColor(0, 0, 0)  # Define a corpreto
        r = dili_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        datata = sindicancia.data_portaria
        dia = f"{datata.day:02d}"
        mes = datata.month
        ano = datata.year

        texto = f'Tendo este Sindicante sido designado por Vossa Senhoria, e por motivo ter  finalizado os trabalhos referentes à Sindicância, instaurada pela Portaria nº {sindicancia.numero}  de {dia} de {mes_escrito(mes)} de {ano}, qual figuram como Sindicados o'
        texto1='estão conclusos  remeto a V.S.ª estes autos com xxxxxxxxxxxxxxxx folhas, para fins de Solução.'
        sindicados_texto = ' '.join([f" {s.posto_sindicado} {s.nome}, RG {s.rgpm} PMMT," for s in sindicados])
        dili_paragraph = doc.add_paragraph()
        dili_run = dili_paragraph.add_run(f"{texto} {sindicados_texto} {texto1}")

        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dili_paragraph.paragraph_format.first_line_indent = Pt(30)

        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor preto
        r = dili_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')



    nome_sindicante(doc,usuario,'Respeitosamente')


    rodape(doc=doc,usuario=usuario)
    nome = f'remessa sind{sindicancia.numero}'

    oficio.save()




    return gera_word(doc,nome)

# remessa dos autos ______________________________________________________________________
@login_required(login_url='/')
def criar_oficio(request, sindicancia_id):



    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    sindicados = sindicancia.sindicados.all()
    testemunhas = sindicancia.testemunhas.all()
    ofendidos = sindicancia.ofendidos.all()
    email = request.user.telefone


    # Adicione as queries para testemunhas, ofendidos e ofícios conforme necessário
    # testemunhas = sindicancia.testemunhas.all()
    # ofendidos = sindicancia.ofendidos.all()
    # oficios = sindicancia.oficios.all()
    vixi= len(sindicados)





    context = {
        'sindicancia': sindicancia,
        'sindicados': sindicados,
        'testemunhas': testemunhas,
        'ofendidos': ofendidos,
        'user_email': email,

        # Adicione as queries para testemunhas, ofendidos e ofícios conforme necessário
        # 'testemunhas': testemunhas,
        # 'ofendidos': ofendidos,

         'vixi': vixi,

    }


    return render(request, 'botoes.html', context)

@login_required(login_url='/')
def cadastrar_notificacao_test(request, sindicancia_id, condicao):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    usuario = request.user


    if request.method == 'POST':
        if condicao == 'test':
            form = NotificarTestForm(request.POST, sindicancia_id=sindicancia_id)
        elif condicao == 'ofen':
            form = NotificarOfenForm(request.POST, sindicancia_id=sindicancia_id)
        elif condicao == 'sind':
            form = NotificarSindForm(request.POST, sindicancia_id=sindicancia_id)
        if form.is_valid():
            oficio=form.save(commit=False)
            oficio.id_portaria=sindicancia.id





            nome_notificado = form.cleaned_data['nome_destinatario']
            data_inquiricao = form.cleaned_data.get('data')  # Já é um objeto datetime.date
            hora_inquiricao = form.cleaned_data.get('hora')
            endereco_testemunha = nome_notificado.endereco





            # Formatar a data e o dia da semana
            dia_inquiricao = data_inquiricao.strftime('%d/%m/%Y')
            dia_semana_ingles = data_inquiricao.strftime('%A')
            dias_da_semana = {
                'Monday': 'segunda-feira',
                'Tuesday': 'terça-feira',
                'Wednesday': 'quarta-feira',
                'Thursday': 'quinta-feira',
                'Friday': 'sexta-feira',
                'Saturday': 'sábado',
                'Sunday': 'domingo'
            }

            # Obtenha o nome do dia da semana em inglês


            # Converta para português usando o dicionário
            dia_semana = dias_da_semana.get(dia_semana_ingles, dia_semana_ingles)



            # Criar o documento .docx
            template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'relatorio_modelo.docx')
            doc = Document(template_path)

            # Preencher o documento com os dados necessários
            cabecalho(doc, usuario)

            criar_paragrafo(doc, '')
            oficio.id_portaria=sindicancia_id
            oficio.numero=numero_oficio(sindicancia_id)

            titulo_oficio(doc,usuario, sindicancia_id)
            criar_paragrafo(doc, '')

            data_original = f'{sindicancia.data_portaria}'
            data_obj = datetime.strptime(data_original, '%Y-%m-%d')
            data_formatada = data_obj.strftime('%d.%m.%y')






            if condicao == 'test':
                condicao = 'Testemunha'
                cargoouendereco=endereco_testemunha
            elif condicao == 'ofen':
                condicao = 'Ofendido'
                cargoouendereco = endereco_testemunha
            else:
                condicao = 'Sindicado'
                cargoouendereco = endereco_testemunha

            assunto = 'Solicitação (FAZ)'

            texto_solicitacao = (
                f'Notifico Vossa Senhoria a comparecer no {usuario.unidade}, {usuario.rua}, nº {usuario.numero},'
                f'Bairro: {usuario.bairro},Cidade: {usuario.cidade}, no dia {data_inquiricao.strftime("%d/%m/%Y")} ({dia_semana}) '
                f'às {hora_inquiricao}, fins de prestar esclarecimentos na condição de {condicao} '
                f'na Sindicância Portaria nº {sindicancia.numero}, de {data_formatada}, em referência a fim de ser inquirido '
                f'sobre os fatos narrados na portaria. Dúvidas entrar em contato por meio do telefone: '
                f'{usuario.telefone} {usuario.posto} {usuario.nome_completo}.'
            )

            criar_paragrafo(doc, f'{nome_notificado}/ {cargoouendereco}')

            criar_paragrafo(doc, f'Ref: Portaria nº{sindicancia.numero}, datada de {data_formatada}.')
            criar_paragrafo(doc, f'Assunto: {assunto}')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, f"{texto_solicitacao}", lado='justificado')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')

            nome_sindicante(doc, usuario, 'Atenciosamente')
            rodape(doc, usuario)

            # Gerar o arquivo para download
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="notificacao_{condicao}-{nome_notificado}.docx"'
            doc.save(response)

            oficio.tipo=f'Notificação de {condicao}'

            oficio.save()
            form.save()
            return response
    else:
        if condicao == 'test':
            form = NotificarTestForm(sindicancia_id=sindicancia_id)
        elif condicao == 'sind':
            form = NotificarSindForm(sindicancia_id=sindicancia_id)
        elif condicao == 'ofen':
            form = NotificarOfenForm(sindicancia_id=sindicancia_id)



    return render(request, 'gerar_notificacao_test.html', {'form': form, 'usuario': usuario})

@login_required(login_url='/')
def Oficio_prazo (request, sindicancia_id, condicao):

    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    usuario = request.user


    if request.method == 'POST':


        if condicao == 'prorrogacao':
            form = PrazoForm(request.POST, sindicancia_id=sindicancia_id)
            condicao='Prorrogação'


        elif condicao == 'dilacao':
            form = PrazoForm(request.POST, sindicancia_id=sindicancia_id)
            condicao = 'Dilação'
        elif condicao == 'sobrestamento':
            form = PrazoForm(request.POST, sindicancia_id=sindicancia_id)
            condicao = 'Sobrestamento'
        if form.is_valid():
            oficio = form.save(commit=False)


            oficio.tipo= f'Solicitação de {condicao}'


            motivo = form.cleaned_data.get('motivo')  # Já é um objeto datetime.date


            # Criar o documento .docx
            template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'relatorio_modelo.docx')
            doc = Document(template_path)

            # Preencher o documento com os dados necessários
            cabecalho(doc, usuario)

            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')
            titulo_oficio(doc, usuario, sindicancia_id)
            criar_paragrafo(doc, '')



            oficio.numero=numero_oficio(sindicancia_id)
            oficio.id_portaria=sindicancia_id
            oficio.nome_destinatario= sindicancia.delegante
            data_original = f'{sindicancia.data_portaria}'
            data_obj = datetime.strptime(data_original, '%Y-%m-%d')
            data_formatada = data_obj.strftime('%d.%m.%y')

            if condicao == 'Prorrogação':

                texto = f'Solicito-vos, nos termos do Art 1º parágrafo 1º da portaria 01/QCG/CORREGPM de 22.01.2018, a Prorrogação de 20(vinte) dias, para a conclusão dos trabalhos, da Sindicância designado por Portaria nº{sindicancia.numero}, datada de {data_formatada}.'


            elif condicao == 'Dilação':
                texto = f'Solicito-vos, nos termos do Art 1º parágrafo 2º da portaria 01/QCG/CORREGPM de 22.01.2018, a Dilação de 30(trinta) dias, para a conclusão dos trabalhos, da Sindicância designado por Portaria nº{sindicancia.numero}, datada de {data_formatada}.'

            elif condicao == 'Sobrestamento':
                texto = f'Solicito-vos, nos termos do Art 1º parágrafo 3º da portaria 01/QCG/CORREGPM de 22.01.2018, o Sobrestamento de 30(trinta) dias, para a conclusão dos trabalhos, da Sindicância designado por Portaria nº{sindicancia.numero}, datada de {data_formatada}.'






            assunto = 'Solicitação (FAZ)'



            #criar_paragrafo(doc, f'{nome_notificado}')

            criar_paragrafo(doc, f'Ao Senhor {sindicancia.posto_delegante} {oficio.nome_destinatario} -')
            criar_paragrafo(doc, f'{sindicancia.funcao_delegante}')
            criar_paragrafo(doc, f'Assunto: {assunto}')
            criar_paragrafo(doc, f'Ref: Portaria nº{sindicancia.numero}, datada de {data_formatada}.')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, 'Senhor Comandante,',iniciar=120)
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, f"{texto}", lado='justificado')
            criar_paragrafo(doc, f'Outrossim, informo que tal solicitação prende-se ao fato de {motivo}.',lado='justificado')

            criar_paragrafo(doc, '')

            nome_sindicante(doc, usuario, 'Respeitosamente')
            rodape(doc, usuario)

            form.save()
            oficio.save()

            # Gerar o arquivo para download
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{condicao}.docx"'
            doc.save(response)

            form.save()
            oficio.save()
            return response
    else:
        if condicao == 'prorrogacao':
            form = PrazoForm(sindicancia_id=sindicancia_id)
        elif condicao == 'dilacao':
            form = PrazoForm(sindicancia_id=sindicancia_id)
        elif condicao == 'sobrestamento':
            form = PrazoForm(sindicancia_id=sindicancia_id)

    return render(request, 'gerar_prazo.html', {'form': form, 'usuario': usuario})



@login_required(login_url='/')
def Oficio_diverso(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    usuario = request.user

    if request.method == 'POST':
        form = Oficio_diversoForm(request.POST, sindicancia_id=sindicancia_id)
        condicao = 'Sobrestamento'
        if form.is_valid():
            oficio = form.save(commit=False)
            oficio.tipo = f'Solicitação de {condicao}'
            oficio.motivo = form.cleaned_data.get('motivo')
            oficio.cargofuncao = form.cleaned_data.get('cargofuncao')

            # Criar o documento .docx
            template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'relatorio_modelo.docx')
            doc = Document(template_path)


            #data_inquiricao = form.cleaned_data.get('data')  # Já é um objeto datetime.date


            # Preencher o documento com os dados necessários
            cabecalho(doc, usuario)
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')
            titulo_oficio(doc, usuario, sindicancia_id)
            criar_paragrafo(doc, '')

            oficio.numero = numero_oficio(sindicancia_id)
            oficio.id_portaria = sindicancia_id
            oficio.nome_destinatario = form.cleaned_data.get('nome_destinatario')
            data_original = f'{sindicancia.data_portaria}'
            data_obj = datetime.strptime(data_original, '%Y-%m-%d')
            data_formatada = data_obj.strftime('%d.%m.%y')

            texto = f'Solicito-vos, nos termos do Art 1º parágrafo 1º da portaria 01/QCG/CORREGPM de 22.01.2018, a Prorrogação de 20(vinte) dias, para a conclusão dos trabalhos, da Sindicância designado por Portaria nº{sindicancia.numero}, datada de {data_formatada}.'

            assunto = 'Solicitação (FAZ)'

            criar_paragrafo(doc, f'A(o) Senhor(a)  {oficio.nome_destinatario} -')
            criar_paragrafo(doc, f'{oficio.cargofuncao}')
            criar_paragrafo(doc, f'Assunto: {assunto}')
            criar_paragrafo(doc, f'Ref: Portaria nº{sindicancia.numero}, datada de {data_formatada}.')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, f'Senhor(a) {oficio.cargofuncao},', iniciar=120)
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, f"{oficio.motivo}", lado='justificado')

            criar_paragrafo(doc, '')

            nome_sindicante(doc, usuario, 'Respeitosamente')
            rodape(doc, usuario)

            form.save()
            oficio.save()

            nome = f'Ofício nº {oficio.numero}'
            return gera_word(doc, nome)
    else:
        # Inicializa o formulário para requisições GET
        form = Oficio_diversoForm(sindicancia_id=sindicancia_id)

    return render(request, 'oficio_diverso.html', {'form': form, 'usuario': usuario})


def Juntada(request, sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    usuario = request.user

    if request.method == 'POST':
        form = JuntadaaForm(request.POST, sindicancia_id=sindicancia_id)

        if form.is_valid():
            oficio = form.save(commit=False)
            oficio.tipo = 'Juntada'
            oficio.motivo = form.cleaned_data.get('motivo')
            oficio.cargofuncao = ''
            oficio.numero=1


            # Criar o documento .docx
            template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'relatorio_modelo.docx')
            doc = Document(template_path)

            cabecalho(doc, usuario)
            criar_paragrafo(doc,'')
            criar_paragrafo(doc, '')


            report_title = doc.add_paragraph("JUNTADA")
            report_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = report_title.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.underline = False
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Adicionar espaçamento duplo antes do título "Relatório"
            report_title_format = report_title.paragraph_format
            report_title_format.space_after = Pt(24)  # Espaço duplo
            report_title_format.space_before = Pt(24)  # Espaço duplo

            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')

            datata = datetime.today()
            dia = datata.day
            mes = datata.month
            ano = datata.year



            texto=(f'{dia_escrito(dia)} dias do mês de {mes_escrito(mes)} do ano de {ano_escrito(ano)}, nesta cidade de {usuario.cidade} , Estado de Mato Grosso, faço juntada aos presentes autos de Sindicância dos documentos de Fls. ........... a .............. , do que para constar;')
            criar_paragrafo(doc, f'{texto}',lado='justificado')
            criar_paragrafo(doc, '')
            criar_paragrafo(doc, '')
            nome_sindicante(doc, usuario, '')
            rodape(doc, usuario)

            form.save()
            oficio.save()

            nome = f'Juntada dia {dia}-{mes}'
            return gera_word(doc, nome)
    else:
        # Inicializa o formulário para requisições GET
        form = JuntadaaForm(sindicancia_id=sindicancia_id)

    return render(request, 'juntada.html', {'form': form, 'usuario': usuario})



@login_required(login_url='/')
def Autuacao(request, sindicancia_id):
    template_path = os.path.join(settings.BASE_DIR, 'inicio/templates', 'autuacao.docx')

    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    sindicados = sindicancia.sindicados.all()
    usuario = request.user
    doc = DocxTemplate(template_path)  # DOCUMENTO  EXEMPLO#


    #dia_inicio = sindicancia.data_inicio

    mes_ini = sindicancia.data_inicio.month
    mesescritoini = mes_escrito(mes_ini)
    diaini = sindicancia.data_inicio.strftime('%d') if sindicancia.data_inicio else ''
    anoini = sindicancia.data_inicio.year

    dia_inicio = f'{diaini} de {mesescritoini} de {anoini}'


    posto_delegante = sindicancia.posto_delegante
    nome_delegante = sindicancia.delegante
    funcao_delegante = sindicancia.funcao_delegante
    portaria = sindicancia.numero

    #datada = sindicancia.data_portaria
    mes_port= sindicancia.data_portaria.month
    mesescritoport=mes_escrito(mes_port)
    diaport=sindicancia.data_portaria.strftime('%d') if sindicancia.data_portaria else ''
    anoport=sindicancia.data_portaria.year


    datada= f'{diaport} de {mesescritoport} de {anoport}'
    cidade = usuario.cidade
    unidade = usuario.unidade
    cr = usuario.cr
    unidade1 = usuario.unidade.upper()
    cr1 = usuario.cr.upper()
    rua = usuario.rua
    bairro = usuario.bairro
    numero = usuario.numero
    cep = usuario.cep
    email = usuario.email_bpm
    telefone= usuario.telefone




    #dia_recebido = sindicancia.dia_recebido
    mes_rec = sindicancia.dia_recebido.month
    mesescritorec = mes_escrito(mes_rec)
    diarec = sindicancia.dia_recebido.strftime('%d') if sindicancia.dia_recebido else ''
    anorec = sindicancia.dia_recebido.year

    dia_recebido = f'{diarec} de {mesescritorec} de {anorec}'



    nome_delegada = sindicancia.delegada
    posto_delegada = sindicancia.posto_delegada
    rg_delegada = sindicancia.rg_delegada
    context = {  # VARIÁ
        "dia_inicio": f"{dia_inicio}",
        "anoini": f"{anoini}",
        "cr": f"{cr}",
        "cr1": f"{cr1}",
        "rua": f"{rua}",
        "bairro": f"{bairro}",
        "numero": f"{numero}",
        "telefone": f"{telefone}",
        "cep": f"{cep}",
        "email": f"{email}",
        "cidade": f"{cidade}",
        "unidade": f"{unidade}",
        "unidade1": f"{unidade1}",
        "posto_delegante": f"{posto_delegante}",
        "nome_delegante": f"{nome_delegante}",
        "funcao_delegante": f"{funcao_delegante}",
        "portaria": f"{portaria}",
        "datada": f"{datada}",
        "dia_recebido": f"{dia_recebido}",
        "nome_delegada": f"{nome_delegada}",
        "posto_delegada": f"{posto_delegada}",
        "rg_delegada": f"{rg_delegada}",

    }
    doc.render(context)

    doc.save("documento_com_tpl.docx")

    doc = Document("documento_com_tpl.docx")


    # Salva o documento final
    sindicado_paragraph = doc.add_paragraph()
    sindicado_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sindicado_paragraph.paragraph_format.first_line_indent = Pt(80)
    sindicado_run = sindicado_paragraph.add_run("SINDICANTE: ")
    sindicado_run.font.name = 'Times New Roman'
    sindicado_run.font.size = Pt(12)
    sindicado_run.bold = True  # Deixar a palavra "Sindicado" em negrito
    r = sindicado_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sindicado_run = sindicado_paragraph.add_run(f"{usuario.nome_completo.upper()} - {usuario.posto.upper()}")
    sindicado_run.font.name = 'Times New Roman'

    sindicado_run.font.size = Pt(12)
    r = sindicado_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')



    if len(sindicados)==1: # verifica se tem um ou mais sindicados pra  colocar no plural se for o caso

        sindicado_paragraph = doc.add_paragraph()
        sindicado_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        sindicado_paragraph.paragraph_format.first_line_indent = Pt(80)
        sindicado_run = sindicado_paragraph.add_run("SINDICADO: ")
        sindicado_run.font.name = 'Times New Roman'
        sindicado_run.font.size = Pt(12)
        sindicado_run.bold = True  # Deixar a palavra "Sindicado" em negrito
        r = sindicado_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        sindicado_paragraph = doc.add_paragraph()
        sindicado_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        sindicado_paragraph.paragraph_format.first_line_indent = Pt(80)
        sindicado_run = sindicado_paragraph.add_run("SINDICADOS: ")
        sindicado_run.font.name = 'Times New Roman'
        sindicado_run.font.size = Pt(12)
        sindicado_run.bold = True  # Deixar a palavra "Sindicado" em negrito
        r = sindicado_run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for index, sindicado in enumerate(sindicados):
        if index == 0:
            sindicado_run = sindicado_paragraph.add_run(f"{sindicado.nome.upper()} - {sindicado.posto_sindicado.upper()}")
            sindicado_run.font.name = 'Times New Roman'

            sindicado_run.font.size = Pt(12)
            r = sindicado_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        else:

            sindicado_nome_paragraph = doc.add_paragraph()
            sindicado_nome_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            sindicado_nome_paragraph.paragraph_format.left_indent = Pt(163)
            sindicado_nome_run = sindicado_nome_paragraph.add_run(f"{sindicado.nome.upper()} - {sindicado.posto_sindicado.upper()}")
            sindicado_nome_run.font.name = 'Times New Roman'
            sindicado_nome_run.font.size = Pt(12)
            r = sindicado_nome_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sindicado_paragraph = doc.add_paragraph('')
    sindicado_paragraph = doc.add_paragraph('')
    sindicado_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sindicado_paragraph.paragraph_format.first_line_indent = Pt(82)
    sindicado_run = sindicado_paragraph.add_run("OBJETO: ")
    sindicado_run.font.name = 'Times New Roman'
    sindicado_run.font.size = Pt(12)
    sindicado_run.bold = True  # Deixar a palavra "OBJETO" em negrito
    r = sindicado_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'{sindicancia.historico}')
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dili_paragraph.paragraph_format.left_indent = Pt(140)
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    dili_paragraph = doc.add_paragraph()
    dili_paragraph = doc.add_paragraph()
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'AUTUAÇÃO')
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.bold = True
    dili_run.italic = True
    dili_paragraph.paragraph_format.left_indent = Pt(100)

    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    r = dili_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sindicado_nome_paragraph.paragraph_format.left_indent = Pt(82)
    sindicado_nome_paragraph.paragraph_format.first_line_indent = Pt(50)
    dia= sindicancia.data_inicio.day
    texto=f'{dia_escrito(dia)} dias do mês de {mes_escrito(mes_ini)} do ano de {anoini}, nesta cidade de {usuario.cidade} Estado de Mato Grosso no Quartel do {usuario.unidade}, autuo a Portaria nº {sindicancia.numero}, de {sindicancia.data_portaria.day} de {mes_escrito(sindicancia.data_portaria.month)} de {sindicancia.data_portaria.year}, e demais documentos que a esta junto me foram entregues. Para constar, eu, {usuario.nome_completo} - {usuario.posto} digitei e assinei este termo. '
    sindicado_nome_run = sindicado_nome_paragraph.add_run(f"{texto}")
    sindicado_nome_run.font.name = 'Times New Roman'
    sindicado_nome_run.font.size = Pt(12)
    r = sindicado_nome_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    sindicado_nome_paragraph.paragraph_format.first_line_indent = Pt(90)
    sindicado_nome_run = sindicado_nome_paragraph.add_run(f"{usuario.nome_completo.upper()}- {usuario.posto.upper()}")
    sindicado_nome_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sindicado_nome_run.bold = True
    sindicado_nome_run.italic = True
    sindicado_nome_run.font.name = 'Times New Roman'
    sindicado_nome_run.font.size = Pt(12)
    r = sindicado_nome_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sindicado_nome_paragraph = doc.add_paragraph()
    sindicado_nome_paragraph.paragraph_format.first_line_indent = Pt(100)
    sindicado_nome_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sindicado_nome_run = sindicado_nome_paragraph.add_run("SINDICANTE")
    sindicado_nome_run.italic = True
    sindicado_nome_run.font.name = 'Times New Roman'
    sindicado_nome_run.font.size = Pt(12)
    r = sindicado_nome_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')




    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename= Autuação {portaria}.docx'
    doc.save(response)

    return response





#______________________funções para construção____________________________________________________________
def criar_paragrafo(doc,texto,antes=1,depois=1,lado=0,iniciar=0):

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'{texto}')
    if lado=='centro':
        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_paragraph.paragraph_format.space_before = Pt(antes)
        dili_paragraph.paragraph_format.space_after = Pt(depois)
        #dili_paragraph.paragraph_format.line_spacing = Pt(12)

    elif lado =='justificado':
        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dili_paragraph.paragraph_format.first_line_indent = Pt(40)
        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_paragraph.paragraph_format.space_before = Pt(antes)
        dili_paragraph.paragraph_format.space_after = Pt(depois)
        # dili_paragraph.paragraph_format.line_spacing = Pt(12)

    else:
        dili_paragraph.paragraph_format.first_line_indent = Pt(iniciar)
        dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dili_run.font.name = 'Times New Roman'
        dili_run.font.size = Pt(12)
        dili_paragraph.paragraph_format.space_before = Pt(antes)
        dili_paragraph.paragraph_format.space_after = Pt(depois)
        # dili_paragraph.paragraph_format.line_spacing = Pt(12)

def cabecalho(doc,usuario):
    cr = usuario.cr.upper()
    unidade = usuario.unidade.upper()

    header_texts = [
        "POLÍCIA MILITAR DO ESTADO DE MATO GROSSO",
        f"{cr}",
        f"{unidade}"
    ]
    for text in header_texts:
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

    # ____________________titulo cabeçalho___________________________________________________
def nome_sindicante(doc,usuario,tratamento):
    paragraph1 = doc.add_paragraph()


    paragraph1.paragraph_format.first_line_indent = Pt(120)
    run1 = paragraph1.add_run(f'{tratamento}')


    paragraph1.paragraph_format.space_before = Pt(15)
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    paragraph1.paragraph_format.space_before = Pt(0)
    paragraph1.paragraph_format.space_after = Pt(20)
    paragraph1.paragraph_format.line_spacing = Pt(18)

    tamanho = len(usuario.nome_completo) + len(usuario.posto) + 6

    dili_paragraph = doc.add_paragraph()

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run('_' * tamanho)
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    dili_paragraph.paragraph_format.space_before = Pt(0)
    dili_paragraph.paragraph_format.space_after = Pt(1)
    dili_paragraph.paragraph_format.line_spacing = Pt(12)

    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'{usuario.nome_completo} - {usuario.posto}')
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
    dili_paragraph.paragraph_format.space_before = Pt(0)
    dili_paragraph.paragraph_format.space_after = Pt(1)
    dili_paragraph.paragraph_format.line_spacing = Pt(12)
    dili_paragraph = doc.add_paragraph()
    dili_run = dili_paragraph.add_run(f'RGPPM {usuario.rgpm} - Sindicante')
    dili_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dili_run.font.name = 'Times New Roman'
    dili_run.font.size = Pt(12)
def rodape(doc,usuario):

    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.add_paragraph()
    footer_run = footer_paragraph.add_run("________________________________________________________________________")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)  # Cor preta
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)

    footer_paragraph = footer.add_paragraph()
    footer_run = footer_paragraph.add_run(f"{usuario.rua}, nº {usuario.numero}, Bairro: {usuario.bairro}, Cidade: {usuario.cidade}")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)  # Cor preta
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)

    footer_paragraph = footer.add_paragraph()
    footer_run = footer_paragraph.add_run(f"Cep {usuario.cep} - Telefone: {usuario.telefone}")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)  # Cor preta
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)

    footer_paragraph = footer.add_paragraph()
    footer_run = footer_paragraph.add_run(f"e-mail: {usuario.email_bpm}")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)  # Cor preta
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)

def gera_word(doc,nome):



    # Configurar a resposta HTTP para download do arquivo
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{nome}.docx"'

    # Salvar o documento modificado na resposta
    try:
        doc.save(response)
    except Exception as e:
        return HttpResponse(f"Erro ao salvar o documento: {e}", status=500)

    return response

def numero_oficio(sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)
    oficio = Oficio()

    max_numero_oficio = Oficio.objects.filter(id_portaria=sindicancia.id).aggregate(Max('numero'))['numero__max']

    # Definir o próximo número de ofício
    if max_numero_oficio is not None:
        novo_numero_oficio = int(max_numero_oficio) + 1
        numero=novo_numero_oficio

    else:
        novo_numero_oficio = 1  # Se não houver ofício, começar com 1
        numero=novo_numero_oficio

    return novo_numero_oficio


def titulo_oficio(doc,usuario,sindicancia_id):
    sindicancia = get_object_or_404(Sindicancia, pk=sindicancia_id)

    numero= numero_oficio(sindicancia_id)



    ano_corrente = datetime.now().year
    dia=f"{datetime.now().day:02d}"
    mes= datetime.now().month

    padrao =sindicancia.padrao_oficio
    criar_paragrafo(doc,f'Ofício nº.{numero}/SIND/{padrao}/{ano_corrente}	       {usuario.cidade}, {dia} de {mes_escrito(mes)} de {ano_corrente}.')






